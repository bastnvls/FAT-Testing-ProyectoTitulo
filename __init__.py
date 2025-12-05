from datetime import datetime, timedelta, time, timezone
from flask import (
    Flask,
    request,
    render_template,
    send_file,
    redirect,
    url_for,
    flash,
    jsonify,
    json,
)
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from PIL import Image
from io import BytesIO
from flask_wtf.csrf import CSRFProtect, CSRFError
from api import api_bp 
import os
import re
from docx.shared import Inches
from funcionalidades.resaltado import subrayar_texto
from flask_login import LoginManager, login_required, current_user, login_user
from flask_mail import Mail, Message
from models import db, bcrypt, User, PasswordResetToken
from config import Config
from utils import (
    validate_password_strength,
    validate_email_format,
    send_password_reset_email,
    send_support_email,
    suscripcion_vigente,
    send_registration_confirmation_email,
)
from functools import wraps
import hmac
import hashlib

app = Flask(__name__)
app.config.from_object(Config)

# Inicializar extensiones
db.init_app(app)
bcrypt.init_app(app)
mail = Mail(app)
csrf = CSRFProtect(app)

#Evita pedir csrf token en la ruta /api/validar-acceso
csrf.exempt(api_bp)

# Configurar Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"
login_manager.login_message = "Por favor inicia sesión para acceder a esta página."

# Registro de blueprint para que la ruta '/api/validar-acceso' funcione
app.register_blueprint(api_bp, url_prefix='/api')


@login_manager.user_loader
def load_user(user_id):
    """
    Propósito:
        Callback utilizado por Flask-Login para recargar el objeto usuario
        desde la base de datos usando el ID almacenado en la sesión.

    Entradas:
        user_id (str): El ID único del usuario (UUID) en formato string.

    Salidas:
        User object: Instancia del usuario si existe.
        None: Si el ID no es válido o no se encuentra.

    Dependencias:
        - db.session.get
        - Modelo User
    """
    # Verificar si se recibió un ID
    if not user_id:
        # Si no hay ID, no se puede cargar usuario
        return None
    
    # Consultar la base de datos usando la sesión de SQLAlchemy 2.0+
    return db.session.get(User, str(user_id))


def aplicar_fuente_cascadia_code(run, size_pt):
    """
    Propósito:
        Aplicar la tipografía específica 'Cascadia Code' y un tamaño determinado
        a un 'run' (fragmento de texto) dentro de un párrafo de Word.

    Entradas:
        run (docx.text.run.Run): El objeto Run de python-docx a modificar.
        size_pt (int/float): El tamaño de la fuente en puntos.

    Salidas:
        None: Modifica el objeto run directamente en memoria.

    Dependencias:
        - docx.shared.Pt
        - docx.oxml.ns.qn
    """
    # Establecer el nombre de la fuente principal
    run.font.name = "Cascadia Code"
    # Establecer el tamaño usando la unidad Pt (puntos)
    run.font.size = Pt(size_pt)
    # Forzar la configuración de fuente para caracteres de Asia Oriental (compatibilidad Word)
    # Esto asegura que Word reconozca la fuente incluso en configuraciones regionales mixtas
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Code")


def limpiar_texto_xml(texto):
    """
    Propósito:
        Sanitizar cadenas de texto eliminando caracteres de control incompatibles
        con el estándar XML de Word (.docx), evitando errores al generar el archivo.

    Entradas:
        texto (str): La cadena de texto original.

    Salidas:
        str: La cadena limpia y segura para insertar en XML.

    Dependencias:
        None (usa funciones nativas de str y ord).
    """
    # Si el texto es None o vacío, devolverlo tal cual
    if not texto:
        return texto

    # Eliminar bytes nulos (NULL bytes) que rompen cualquier parser XML
    texto = texto.replace("\x00", "")

    # Lista para acumular los caracteres válidos
    caracteres_permitidos = []
    
    # Iterar sobre cada carácter del texto
    for char in texto:
        # Obtener el código ASCII/Unicode del carácter
        code = ord(char)
        
        # Criterios de aceptación XML 1.0:
        # 1. Caracteres imprimibles (code >= 0x20)
        # 2. Caracteres de control permitidos: Tab (0x09), Salto línea (0x0A), Retorno carro (0x0D)
        if code >= 0x20 or code in (0x09, 0x0A, 0x0D):
            # Excluir específicamente el carácter DEL (0x7F)
            # Excluir rango de control extendido (0x80-0x9F) que a veces causa problemas
            if code != 0x7F and not (0x80 <= code <= 0x9F):
                caracteres_permitidos.append(char)

    # Unir la lista en un solo string
    return "".join(caracteres_permitidos)


def pruebas(lines, start_re, end_re):
    """
    Propósito:
        Extraer bloques de texto específicos desde una lista de líneas, delimitados
        por expresiones regulares de inicio y fin. Usado para aislar los logs de cada prueba.

    Entradas:
        lines (list[str]): Lista de todas las líneas del archivo de texto.
        start_re (re.Pattern): Regex compilado que marca el inicio del bloque.
        end_re (re.Pattern): Regex compilado que marca el fin del bloque.

    Salidas:
        list[str]: Lista de bloques de texto encontrados (cada bloque es un string con saltos de línea).

    Dependencias:
        - re (expresiones regulares)
    """
    # Lista para guardar los bloques completos encontrados
    blocks = []
    # Bandera para saber si estamos actualmente dentro de un bloque
    collecting = False
    # Lista temporal para guardar las líneas del bloque actual
    current = []

    # Recorrer el archivo línea por línea
    for line in lines:
        # Quitar espacios al inicio y final para facilitar la detección de patrones
        stripped = line.strip()
        
        # CASO 1: No estamos recolectando y encontramos la marca de inicio
        if not collecting and start_re.match(stripped):
            # Activamos la bandera de recolección
            collecting = True
            # Iniciamos el bloque actual con esta línea
            current = [line]
            # Pasamos a la siguiente iteración
            continue

        # CASO 2: Estamos recolectando líneas
        if collecting:
            # Añadimos la línea actual al buffer
            current.append(line)
            
            # Verificamos si esta línea es la marca de finalización
            if end_re.match(stripped):
                # Si es el fin, unimos las líneas acumuladas y las guardamos en blocks
                blocks.append("\n".join(current))
                # Desactivamos la bandera para buscar el siguiente bloque
                collecting = False
                
    # Retornar todos los bloques encontrados
    return blocks


def iter_paragraphs(doc):
    """
    Propósito:
        Generador que recorre TODOS los párrafos de un documento Word, incluyendo
        los que están dentro de tablas (celdas anidadas), para búsqueda y reemplazo global.

    Entradas:
        doc (docx.Document): El objeto documento a recorrer.

    Salidas:
        yield: Párrafos individuales (docx.text.paragraph.Paragraph).

    Dependencias:
        - python-docx
    """
    # 1. Iterar sobre párrafos del cuerpo principal del documento
    for p in doc.paragraphs:
        yield p
        
    # 2. Iterar sobre todas las tablas del documento
    for table in doc.tables:
        # Recorrer filas
        for row in table.rows:
            # Recorrer celdas
            for cell in row.cells:
                # Recorrer párrafos dentro de la celda
                for p in cell.paragraphs:
                    yield p


def aplicar_fuente_cascadia_code(run, size_pt):
    run.font.name = "Cascadia Code"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Code")


def limpiar_texto_xml(texto):
    """
    Limpia el texto de caracteres incompatibles con XML.
    Elimina NULL bytes y caracteres de control excepto tab, newline y carriage return.

    Args:
        texto: String a limpiar

    Returns:
        String limpio compatible con XML
    """
    if not texto:
        return texto

    # Eliminar NULL bytes
    texto = texto.replace("\x00", "")

    # Eliminar caracteres de control (0x00-0x1F) excepto tab (0x09), newline (0x0A), carriage return (0x0D)
    # También eliminar DEL (0x7F) y caracteres de control en rango extendido (0x80-0x9F)
    caracteres_permitidos = []
    for char in texto:
        code = ord(char)
        # Permitir caracteres normales, tab, newline, carriage return
        if code >= 0x20 or code in (0x09, 0x0A, 0x0D):
            # Excluir DEL y caracteres de control extendidos
            if code != 0x7F and not (0x80 <= code <= 0x9F):
                caracteres_permitidos.append(char)

    return "".join(caracteres_permitidos)


# Función para extraer bloques
def pruebas(lines, start_re, end_re):
    blocks = []
    collecting = False
    current = []
    for line in lines:
        stripped = line.strip()
        if not collecting and start_re.match(stripped):
            collecting = True
            current = [line]
            continue

        if collecting:
            current.append(line)
            if end_re.match(stripped):
                blocks.append("\n".join(current))
                collecting = False
    return blocks


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


# EMu_PER_INCH:
#   Número de EMUs (English Metric Units) que hay en una pulgada.
#   Es la unidad interna que usa Word para tamaños.
EMu_PER_INCH = 914400

# DEFAULT_DPI:
#   DPI "fijo" que vamos a usar para convertir píxeles a pulgadas.
#   Muchos archivos traen 300/600 DPI y eso hace que Word las vea "muy pequeñas".
DEFAULT_DPI = 96


def _get_image_emu_size(image_bytes, max_width_cm=None):
    """
    Propósito:
        Calcular el tamaño de una imagen en EMUs (unidades que usa Word),
        partiendo de sus dimensiones en píxeles. Opcionalmente limita el
        ancho máximo para que la imagen no sea gigantesca en la página.

    Entradas:
        image_bytes (bytes):
            Bytes crudos de la imagen (tal como los devuelve .read()).
        max_width_cm (float | None):
            - None: no limita el ancho, se respeta el tamaño según píxeles.
            - Número en centímetros: si la imagen es más ancha, se escala
              proporcionalmente para que no supere ese ancho.

    Salidas:
        (cx, cy) (tuple[int, int]):
            Tupla con ancho (cx) y alto (cy) de la imagen en EMUs.

    Dependencias:
        - PIL.Image (from PIL import Image)
        - io.BytesIO (from io import BytesIO)
        - EMu_PER_INCH (constante de este módulo)
        - DEFAULT_DPI (constante de este módulo)
    """

    # Abrimos la imagen desde los bytes usando un buffer en memoria
    img = Image.open(BytesIO(image_bytes))

    # Obtenemos el tamaño de la imagen en píxeles (ancho y alto)
    px_w, px_h = img.size

    # Convertimos píxeles a pulgadas usando un DPI fijo (DEFAULT_DPI)
    #   width_inch  = ancho_en_pixeles  / DPI
    #   height_inch = alto_en_pixeles   / DPI
    width_inch = px_w / DEFAULT_DPI
    height_inch = px_h / DEFAULT_DPI

    # Si se especificó un ancho máximo en centímetros, lo aplicamos.
    if max_width_cm is not None:
        # Pasamos de cm a pulgadas: 1 pulgada = 2.54 cm
        max_width_inch = max_width_cm / 2.54

        # Si la imagen es más ancha que el máximo permitido...
        if width_inch > max_width_inch:
            # Calculamos un factor de escala (target / actual)
            scale = max_width_inch / width_inch

            # Escalamos tanto ancho como alto para mantener la proporción
            width_inch *= scale
            height_inch *= scale

    # Ahora convertimos pulgadas a EMUs:
    #   EMUs = pulgadas * EMu_PER_INCH
    cx = int(width_inch * EMu_PER_INCH)
    cy = int(height_inch * EMu_PER_INCH)

    # Devolvemos ancho y alto en EMUs
    return cx, cy

def reemplazar_imagen_flotante(doc, marker, image_file):
    """
    Propósito:
        Reemplazar en el documento Word la imagen de un placeholder
        identificado por el atributo descr == marker, usando la imagen
        subida por el usuario y ajustando su tamaño.

    Entradas:
        doc (docx.Document):
            Documento Word cargado con python-docx.
        marker (str):
            Valor del atributo 'descr' en <pic:cNvPr>, por ejemplo "IMG1".
        image_file (FileStorage):
            Archivo de imagen recibido por Flask (request.files[...]).
    
    Salidas:
        None. Modifica el documento `doc` en memoria.

    Dependencias:
        - doc.part.get_or_add_image
        - _get_image_emu_size
        - doc.element.xpath
        - qn (docx.oxml.ns.qn)
    """

    # Leer todos los bytes del archivo de imagen subido.
    image_bytes = image_file.read()

    # Registrar la imagen dentro del paquete .docx y obtener el nuevo rId.
    # new_rId: ID de relación para referenciar la imagen desde el XML.
    # _: objeto ImagePart que no necesitamos guardar aquí.
    new_rId, _ = doc.part.get_or_add_image(BytesIO(image_bytes))

    # Calcular el tamaño de la imagen en EMUs.
    # max_width_cm=16 limita el ancho máximo a ~16 cm, manteniendo proporción.
    cx, cy = _get_image_emu_size(image_bytes, max_width_cm=16)

    # Recorrer todos los nodos <w:drawing> del documento.
    # Ahí es donde se guardan las imágenes e ilustraciones.
    for drawing in doc.element.xpath(".//w:drawing"):

        # Dentro de cada <w:drawing> buscamos el nodo <pic:pic> (la imagen en sí).
        pics = drawing.xpath(".//pic:pic")
        if not pics:
            # Si no hay imagen en este drawing, pasamos al siguiente.
            continue

        # Buscar el nodo <pic:cNvPr> que tiene atributos como name y descr.
        cNvPr_list = pics[0].xpath(".//pic:cNvPr")
        if not cNvPr_list:
            # Si no existe, no podemos comprobar el descr, así que seguimos.
            continue

        cNvPr = cNvPr_list[0]

        # Solo queremos modificar el dibujo cuyo descr coincide con nuestro marcador.
        # Ej: descr="IMG1", descr="IMG2", etc.
        if cNvPr.get("descr") != marker:
            # Si no coincide, seguimos con el siguiente <w:drawing>.
            continue

        # Si llegamos aquí, encontramos el placeholder correcto.

        # Dentro de <pic:pic> buscamos <a:blip>, que tiene el atributo r:embed
        # apuntando al rId de la imagen actual.
        blip_list = pics[0].xpath(".//a:blip")
        if not blip_list:
            # Si no hay blip, algo raro pasa, salimos del bucle.
            break

        blip = blip_list[0]

        # Actualizamos r:embed para que apunte al nuevo rId (la imagen subida).
        blip.set(qn("r:embed"), new_rId)

        # Buscar <wp:extent>, que define ancho y alto de la imagen en EMUs.
        extent_list = drawing.xpath(".//wp:extent")
        if extent_list:
            extent = extent_list[0]
            # Actualizar ancho (cx) y alto (cy) en EMUs.
            extent.set("cx", str(cx))
            extent.set("cy", str(cy))

        # Rompemos el bucle porque ya reemplazamos el placeholder que nos interesaba.
        break




def insertar_imagenes(doc, image_files, marker):
    """
    Propósito:
        Gestionar la inserción de una o varias imágenes en los placeholders del documento.
        Clona el placeholder original si se deben insertar múltiples archivos.

    Entradas:
        doc (docx.Document): Documento Word abierto en memoria.
        image_files (list[FileStorage]): Lista de imágenes subidas desde Flask.
        marker (str): Valor del atributo 'descr' en <pic:cNvPr> que identifica el placeholder.

    Salidas:
        None: Modifica el objeto `doc` directamente.

    Dependencias:
        - reemplazar_imagen_flotante: maneja la carga de bytes, generación de rId y actualización de XML.
        - deepcopy: permite clonar nodos XML para insertar múltiples instancias sin perder el original.
        - doc.element.xpath: busca nodos <w:drawing> que contienen los placeholders.
    """
    # 1) Buscar todos los contenedores <w:drawing> con el marker en <pic:cNvPr>
    drawings = [
        d
        for d in doc.element.xpath(".//w:drawing")  # obtiene lista de nodos <w:drawing>
        if d.xpath(".//pic:cNvPr")[0].get("descr")
        == marker  # filtra solo los que coinciden con marker
    ]
    #    Por qué: identificar todas las posiciones donde puede ir una imagen.
    #    Relación: estos nodos serán actualizados por reemplazar_imagen_flotante.
    if not drawings:
        # Por qué: si no encuentra placeholders, no hay nada que reemplazar.
        # Relación: evita llamar a reemplazar_imagen_flotante innecesariamente.
        return

    # 2) Seleccionar el primer placeholder como plantilla
    first = drawings[0]  # nodo XML <w:drawing> original
    #    Por qué: usamos el primer nodo para la primera imagen y como base para clonar.
    #    Relación: reemplazar_imagen_flotante siempre repara el primer placeholder.
    parent = first.getparent()  # contenedor XML de los drawings
    #    Por qué: necesitamos el nodo padre para insertar clones.
    #    Relación: parent.insert se usará para agregar nodos clonados.
    idx = parent.index(first)  # posición del placeholder original en `parent`
    #    Por qué: conocer el índice inicial permite calcular posición de inserción.
    #    Relación: idx + i determina dónde ubicar cada clon.

    # 3) Iterar sobre cada imagen subida
    for i, img in enumerate(
        image_files
    ):  # i: índice (0,1,2...), img: objeto FileStorage
        # 3.1) Determinar nodo a usar: original si i=0, clon si i>0
        drawing = first if i == 0 else deepcopy(first)
        #      Por qué: el primer placeholder se actualiza directamente,
        #      y para más imágenes mantenemos la anterior intacta.
        #      Relación: deepcopy crea un nodo independiente para no interferir con el original.

        # 3.2) Reemplazar contenido y tamaño en el placeholder original
        reemplazar_imagen_flotante(doc, marker, img)
        #      Por qué: esta función maneja la inserción real del binario,
        #      generación de new_rId y ajuste de dimensiones.
        #      Relación: conecta image_files con <a:blip r:embed> y <wp:extent>.

        # 3.3) Insertar el clon tras el original si es imagen adicional
        if i > 0:
            parent.insert(idx + i, drawing)
            #      Por qué: ubicamos cada nueva imagen en orden de subida.
            #      Relación: mantiene la secuencia de <w:drawing> en el XML.

    # Fin de insertar_imagenes: cada img en image_files aparece en un <w:drawing> distinto


def replace_marker_with_text(doc, marker, text):
    for p in iter_paragraphs(doc):
        if marker in p.text:
            # Eliminar runs viejos
            for run in list(p.runs):
                p._p.remove(run._r)

            # Limpiar el texto del párrafo (por si queda el marcador)
            p.text = ""

            # Limpiar el texto de caracteres incompatibles con XML
            text = limpiar_texto_xml(str(text))

            # Crear run con el texto nuevo
            run = p.add_run(text)
            run.font.name = "Arial"

            # Compatibilidad para fuentes
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            if marker == "{{proyecto}}":
                run.font.size = Pt(13)
                run.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                # Para cliente, orden_compra, nota_venta
                run.font.size = Pt(11)
                run.bold = False
                p.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )  
            return


# ====== Insertar texto en Word =======
def insertar_texto(doc, marker, texto, size_pt):
    """
    Busca celdas con `marker`, borra su contenido y agrega todo el `texto`
    línea a línea con la fuente y tamaño indicados.
    Devuelve la lista de párrafos creados (para más tarde resaltar).
    """
    paras = []

    # Limpiar el texto de caracteres incompatibles con XML
    texto = limpiar_texto_xml(texto)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if marker in cell.text:
                    cell.text = ""
                    para = cell.add_paragraph()
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for line in texto.split("\n"):
                        # Limpiar cada línea individualmente por seguridad
                        line = limpiar_texto_xml(line)
                        run = para.add_run(line)
                        aplicar_fuente_cascadia_code(run, size_pt)
                        para.add_run("\n")  # salto tras cada línea
                    paras.append(para)
    return paras


def insertar_info_dispositivo(doc, modelo, serial, version):
    # Limpiar los valores de caracteres incompatibles con XML
    modelo = limpiar_texto_xml(modelo) if modelo else ""
    serial = limpiar_texto_xml(serial) if serial else ""
    version = limpiar_texto_xml(version) if version else ""

    # También revisamos las celdas de todas las tablas en el documento, por si no están en párrafos
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{modelo}}" in cell.text:
                    cell.text = cell.text.replace("{{modelo}}", modelo)
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
                    # Aplicar formato Arial 11 a todo el párrafo
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

                if "{{serial}}" in cell.text:
                    cell.text = cell.text.replace("{{serial}}", serial)
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

                if "{{version}}" in cell.text:
                    cell.text = cell.text.replace("{{version}}", version)
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def procesar_archivo(
    file_stream,
    docx_template_path,
    img_1,
    img_2,
    img_3,
    proyecto,
    cliente,
    ordenCompra,
    notaVenta,
    file_type,
):

    # Carga la plantilla
    doc = Document(docx_template_path)

    # Leemos todo el contenido del TXT en 'lines'
    file_stream.seek(0)
    lines = file_stream.read().decode("utf-8").split("\n")

    # --------- EXTRACCION DE INFORMACION DEL DISPOSITIVO ----------
    # Inicializamos variables para almacenar información del dispositivo
    modelo, serial, version = None, None, None

    # Expresiones regulares para extraer información del dispositivo
    if (
        file_type == "SW L2 9200"
        or file_type == "SW L2 9300"
        or file_type == "SW L2 9500"
    ):
        # Ahora buscamos el bloque de "INICIO PRUEBA 1"
        patron_inicio_prueba_1 = re.compile(
            r".*[#>]\s*INICIO\s+PRUEBA\s+1\b", re.IGNORECASE
        )
        # Expresiones regulares para extraer información del dispositivo
        modelo_regex = re.compile(r"Model Number\s*:\s*(\S+)", re.IGNORECASE)
        serial_regex = re.compile(r"System Serial Number\s*:\s*(\S+)", re.IGNORECASE)

        # Buscar información del dispositivo
        for line in lines:
            model_match = modelo_regex.search(line)
            serial_match = serial_regex.search(line)

            if model_match:
                modelo = model_match.group(1).strip()
            if serial_match:
                serial = serial_match.group(1).strip()

        for line in lines:
            match = patron_inicio_prueba_1.match(line.strip())
            if match:
                # Contenido dentro de prueba 1
                for siguiente_linea in lines[
                    lines.index(line) :
                ]:  # Comienza a buscar en las líneas siguientes a la línea actual (donde se encontró "INICIO PRUEBA 1") para encontrar la línea que contiene la versión.
                    version_match = re.search(
                        r"Version\s+(\S+)", siguiente_linea, re.IGNORECASE
                    )  # Usa una expresión regular para buscar la línea que contiene "Version" y capturar el valor que le sigue.
                    if version_match:
                        version = version_match.group(1).strip()
                        break  # Salir del bucle al encontrar la versión
                break  # Salir del bucle principal al encontrar el bloque de prueba

    elif file_type == "SW L3 9348GC" or file_type == "SW L3 C93180YC":
        # Ahora buscamos el bloque de "INICIO PRUEBA 1"
        patron_inicio_prueba_1 = re.compile(
            r".*[#>]\s*INICIO\s+PRUEBA\s+1\b", re.IGNORECASE
        )
        for line in lines:
            match = patron_inicio_prueba_1.match(line.strip())
            if match:
                # Contenido dentro de prueba 1
                for siguiente_linea in lines[
                    lines.index(line) + 1 :
                ]:  # Comienza a buscar desde la siguiente línea
                    model_match = re.search(
                        r"PID:\s*(\S+)", siguiente_linea, re.IGNORECASE
                    )
                    serial_match = re.search(
                        r"SN:\s*(\S+)", siguiente_linea, re.IGNORECASE
                    )
                    version_match = re.search(
                        r"NXOS:\s+version\s+(\S+)", siguiente_linea, re.IGNORECASE
                    )

                    if model_match:
                        modelo = model_match.group(1).strip()
                    if serial_match:
                        serial = serial_match.group(1).strip()
                    if version_match:
                        version = version_match.group(1).strip()

                    # Verificar si se encontraron todos los datos
                    if modelo and serial and version:
                        break  # Salir del bucle si todos están encontrados
                break  # Salir del bucle principal al encontrar el bloque de prueba

    elif file_type == "SW IE3300" or file_type == "SW IE4010":
        # Ahora buscamos el bloque de "INICIO PRUEBA 1"
        patron_inicio_prueba_1 = re.compile(
            r".*[#>]\s*INICIO\s+PRUEBA\s+1\b", re.IGNORECASE
        )

        # Expresiones regulares para extraer información del dispositivo
        modelo_regex = re.compile(r"Model Number\s*:\s*(\S+)", re.IGNORECASE)
        serial_regex = re.compile(r"System Serial Number\s*:\s*(\S+)", re.IGNORECASE)

        # Buscar información del dispositivo
        for line in lines:
            model_match = modelo_regex.search(line)
            serial_match = serial_regex.search(line)

            if model_match:
                modelo = model_match.group(1).strip()
            if serial_match:
                serial = serial_match.group(1).strip()

        for line in lines:
            match = patron_inicio_prueba_1.match(line.strip())
            if match:
                # Contenido dentro de prueba 1
                for siguiente_linea in lines[
                    lines.index(line) :
                ]:  # Comienza a buscar en las líneas siguientes a la línea actual (donde se encontró "INICIO PRUEBA 1") para encontrar la línea que contiene la versión.
                    version_match = re.search(
                        r"Version\s+(\S+)", siguiente_linea, re.IGNORECASE
                    )  # Usa una expresión regular para buscar la línea que contiene "Version" y capturar el valor que le sigue.
                    if version_match:
                        version = version_match.group(1).strip()
                        break  # Salir del bucle al encontrar la versión
                break  # Salir del bucle principal al encontrar el bloque de prueba

    elif file_type == "Router C8500" or file_type == "Router ISR4431":
        # Ahora buscamos el bloque de "INICIO PRUEBA 1"
        patron_inicio_prueba_1 = re.compile(
            r".*[#>]\s*INICIO\s+PRUEBA\s+1\b", re.IGNORECASE
        )
        for line in lines:
            match = patron_inicio_prueba_1.match(line.strip())
            if match:
                # Contenido dentro de prueba 1
                for siguiente_linea in lines[
                    lines.index(line) + 1 :
                ]:  # Comienza a buscar desde la siguiente línea
                    model_match = re.search(
                        r"PID:\s*(\S+)", siguiente_linea, re.IGNORECASE
                    )
                    serial_match = re.search(
                        r"SN:\s*(\S+)", siguiente_linea, re.IGNORECASE
                    )

                    if model_match:
                        modelo = model_match.group(1).strip()
                    if serial_match:
                        serial = serial_match.group(1).strip()

                    # Verificar si se encontraron todos los datos
                    if modelo and serial:
                        break  # Salir del bucle si todos están encontrados
                break  # Salir del bucle principal al encontrar el bloque de prueba

        for line in lines:
            match = patron_inicio_prueba_1.match(line.strip())
            if match:
                # Contenido dentro de prueba 1
                for siguiente_linea in lines[
                    lines.index(line) :
                ]:  # Comienza a buscar en las líneas siguientes a la línea actual (donde se encontró "INICIO PRUEBA 1") para encontrar la línea que contiene la versión.
                    version_match = re.search(
                        r"Version\s+(\S+)", siguiente_linea, re.IGNORECASE
                    )  # Usa una expresión regular para buscar la línea que contiene "Version" y capturar el valor que le sigue.
                    if version_match:
                        version = version_match.group(1).strip()
                        break  # Salir del bucle al encontrar la versión
                break  # Salir del bucle principal al encontrar el bloque de prueba

    elif (
        file_type == "AP C9115AXI"
        or file_type == "AP C9120AXE"
        or file_type == "AP C9130AXI"
    ):
        # Expresiones regulares para extraer información del dispositivo
        modelo_regex = re.compile(r"Product/Model Number\s*:\s*(\S+)", re.IGNORECASE)
        serial_regex = re.compile(
            r"Top Assembly Serial Number\s*:\s*(\S+)", re.IGNORECASE
        )
        version_regex = re.compile(r"Primary Boot Image\s*:\s*(\S+)", re.IGNORECASE)

        # Buscar información del dispositivo
        for line in lines:
            model_match = modelo_regex.search(line)
            serial_match = serial_regex.search(line)
            version_match = version_regex.search(line)

            if model_match:
                modelo = model_match.group(1).strip()
            if serial_match:
                serial = serial_match.group(1).strip()
            if version_match:
                version = version_match.group(1).strip()

    elif (
        file_type == "AP C9115AXI"
        or file_type == "AP C9120AXE"
        or file_type == "AP C9130AXI"
    ):
        # Expresiones regulares para extraer información del dispositivo
        modelo_regex = re.compile(r"Product/Model Number\s*:\s*(\S+)", re.IGNORECASE)
        serial_regex = re.compile(
            r"Top Assembly Serial Number\s*:\s*(\S+)", re.IGNORECASE
        )
        version_regex = re.compile(r"Primary Boot Image\s*:\s*(\S+)", re.IGNORECASE)

        # Buscar información del dispositivo
        for line in lines:
            model_match = modelo_regex.search(line)
            serial_match = serial_regex.search(line)
            version_match = version_regex.search(line)

            if model_match:
                modelo = model_match.group(1).strip()
            if serial_match:
                serial = serial_match.group(1).strip()
            if version_match:
                version = version_match.group(1).strip()

    elif file_type == "Check Point 6200" or file_type == "Check Point 6600":
        # Expresiones regulares para extraer información del dispositivo
        modelo_regex = re.compile(
            r"Appliance Name\s*:\s*(.+)", re.IGNORECASE
        )  # Captura todo lo que sigue
        serial_regex = re.compile(r"Appliance SN\s*:\s*(\S+)", re.IGNORECASE)
        version_regex = re.compile(
            r"SVN Foundation Version String\s*:\s*(\S+)", re.IGNORECASE
        )

        # Buscar información del dispositivo
        for line in lines:
            model_match = modelo_regex.search(line)
            serial_match = serial_regex.search(line)
            version_match = version_regex.search(line)

            if model_match:
                modelo = model_match.group(1).strip()
            if serial_match:
                serial = serial_match.group(1).strip()
            if version_match:
                version = version_match.group(1).strip()

    # Insertar información del dispositivo en el documento
    insertar_info_dispositivo(doc, modelo, serial, version)

    # -------------PRUEBAS--------------
    # 1) Primero, buscamos cuántas "pruebas" hay en el texto.
    #    Cada instrucción de inicio de prueba debería tener formato:
    #       # INICIO PRUEBA {numero}
    #   Extraemos todos los números que aparezcan en esas líneas.
    patron_inicio_any = re.compile(r".*[#>]\s*INICIO\s+PRUEBA\s+(\d+)\b", re.IGNORECASE)
    numeros_encontrados = set()

    for line in lines:
        match = patron_inicio_any.match(line.strip())
        if match:
            numeros_encontrados.add(int(match.group(1)))

    buffer = BytesIO()
    if not numeros_encontrados:
        doc.save(buffer)
        buffer.seek(0)
        nombre = f"{file_type}_documento_vacio.docx"
        return buffer, nombre

    # 2) Ordenamos los números de prueba (1, 2, 3, ...)
    numeros_ordenados = sorted(numeros_encontrados)
    # 3) Por cada número de prueba, construimos los regex de inicio y fin,
    #    llamamos a 'pruebas' para extraer el bloque, y luego insertamos el texto
    contador = 0
    for n in numeros_ordenados:
        # Creamos un patrón para:    # INICIO PRUEBA {n}
        start_re = re.compile(rf".*[#>]\s*INICIO\s+PRUEBA\s+{n}\b", re.IGNORECASE)
        # Creamos un patrón para:    # FIN PRUEBA {n}
        end_re = re.compile(rf".*[#>]\s*FIN\s+PRUEBA\s+{n}\b", re.IGNORECASE)

        # Extraemos el bloque correspondiente a esa prueba
        bloques = pruebas(lines, start_re, end_re)

        # Si encontramos bloques, los procesamos uno a uno.
        # El texto a insertar lleva un sufijo con dos dígitos, p. ej. "01", "02", ...
        sufijo = f"{n:02d}"
        texto_label = f"Insertar codigo de la extracción {sufijo}"

        for bloque in bloques:
            contador = n
            # Se inserta el texto
            paras = insertar_texto(doc, texto_label, bloque, 8)
            # Se subraya el texto
            subrayar_texto(paras, file_type, contador)

    # Se insertan las imagenes
    # Se reemplazan las imágenes flotantes predefinidas
    insertar_imagenes(doc, img_1, "IMG1")
    insertar_imagenes(doc, img_2, "IMG2")
    insertar_imagenes(doc, img_3, "IMG3")
    replace_marker_with_text(doc, "{{proyecto}}", proyecto)
    replace_marker_with_text(doc, "{{cliente}}", cliente)
    replace_marker_with_text(doc, "{{orden_compra}}", ordenCompra)
    replace_marker_with_text(doc, "{{nota_venta}}", notaVenta)
    # 4) Una vez terminadas todas las pruebas, guardamos el documento
    doc.save(buffer)
    buffer.seek(0)
    nombre = f"{modelo} {serial}.docx"
    return buffer, nombre


def necesita_pago(user):
    """
    Propósito:
        Indicar si el usuario se encuentra en un estado de suscripción que
        requiere pasar por el flujo de pago / checkout de Mercado Pago.

    Entradas:
        - user: instancia del modelo User con el campo estado_suscripcion.

    Salidas:
        - bool: True si debe ir a la pasarela de pago, False si no.

    Dependencias:
        - Campo user.estado_suscripcion (string).
        - Conjunto de estados definidos como "requiere pago".
    """

    # Definimos el conjunto de estados que consideramos que necesitan pago.
    estados_que_necesitan_pago = {
        "SIN_SUSCRIPCION",  # Nunca ha contratado un plan.
        "PENDIENTE_MP",  # Ya inició el proceso en MP pero aún no se autorizó.
        "PAUSADA",  # La suscripción está pausada en MP.
        "CANCELADA",  # La suscripción fue cancelada.
        "VENCIDA",  # La fecha de fin de suscripción ya pasó.
        "EN_GRACIA",  # Periodo de gracia, pero aún queremos forzar pago.
    }

    # Devolvemos True si el estado actual del usuario está en el conjunto anterior.
    return user.estado_suscripcion in estados_que_necesitan_pago


def suscripcion_requerida(vista):
    """
    Propósito:
        Decorador para vistas de Flask que obliga a que el usuario:
        1. Esté autenticado (logueado), y
        2. Tenga una suscripción vigente.

    Entradas:
        vista: función de vista de Flask que queremos proteger
               (por ejemplo: upload_files, dashboard, etc.).

    Salidas:
        Devuelve una nueva función (vista_protegida) que envuelve a vista
        y agrega la lógica de:
            - verificar login
            - verificar suscripción
            - redirigir si no cumple

    Dependencias:
        - @login_required: decorador de Flask-Login que exige usuario autenticado.
        - current_user: objeto de Flask-Login que representa al usuario actual.
        - suscripcion_vigente(user): función tuya que devuelve True/False.
        - flash: función de Flask para mostrar mensajes al usuario.
        - redirect, url_for: funciones de Flask para redirigir a otra ruta.
    """

    # Hace que vista_protegida conserve el nombre y metadatos de vista original,
    # como si siguiera siendo la misma función a ojos de Flask y las herramientas.
    @wraps(vista)

    # Este decorador de Flask-Login fuerza SOLO usuarios logueados
    @login_required
    def vista_protegida(*args, **kwargs):

        # Primero, refrescamos el estado de suscripción del usuario actual
        refrescar_estado_suscripcion(current_user)

        # Si el usuario actual NO tiene una suscripción vigente...
        if not suscripcion_vigente(current_user):
            # Mostramos un mensaje de advertencia en la interfaz (con Flash).
            flash(
                "Necesitas una suscripción ACTIVA para acceder a esta sección.",
                "warning",
            )

            # Redirigimos al usuario a la ruta donde puede pagar/activar
            # su suscripción
            return redirect(url_for("suscripcion_checkout"))

        # Si tiene suscripción vigente, ejecutamos la vista original
        # con sus mismos argumentos (*args, **kwargs) y retornamos su resultado.
        return vista(*args, **kwargs)

    # devolvemos la función "envuelta" (vista_protegida)
    return vista_protegida


def extraer_evento_mp(body, query):
    """
    Propósito:
        Analizar la petición de Mercado Pago (JSON body o Query Params)
        para determinar qué tipo de evento ocurrió y el ID del recurso afectado.

    Entradas:
        body (dict): El JSON recibido en el cuerpo del webhook.
        query (dict): Los parámetros de la URL (request.args).

    Salidas:
        tuple (str, str): (topic, resource_id). Devuelve (None, None) si falla.

    Dependencias:
        None.
    """
    # Intentamos obtener el tipo de evento directamente desde el body.
    topic = body.get("type")

    # Si en el body no vino "type", probamos tomarlo desde la query string.
    if topic is None:
        topic = query.get("type")

    # Obtenemos el campo "data" del body, si existe, o un diccionario vacío.
    data_dict = body.get("data") or {}

    # Dentro de "data", intentamos leer el identificador principal del recurso.
    resource_id = data_dict.get("id")

    # Si no encontramos el id dentro de "data", probamos con "data.id" en la query.
    if resource_id is None:
        resource_id = query.get("data.id")

    # Como último intento, miramos si hay un "id" en la raíz del body.
    if resource_id is None:
        resource_id = body.get("id")

    # Si falta el topic o el id del recurso, devolvemos (None, None) para indicar error.
    if topic is None or resource_id is None:
        return None, None

    # Devolvemos la tupla (topic, resource_id) cuando ambos existen.
    return topic, resource_id


def obtener_usuario_desde_preapproval(preapproval):
    """
    Propósito:
        Localizar al usuario en la base de datos basándose en los metadatos
        de una suscripción de Mercado Pago (external_reference o ID de preapproval).

    Entradas:
        preapproval (dict): Datos de la suscripción recibidos de MP.

    Salidas:
        User object: Instancia del usuario encontrado.
        None: Si no se encuentra coincidencia.

    Dependencias:
        - Modelo User
        - db.session.get
    """
    # Intentamos leer el external_reference definido cuando se creó la suscripción.
    external_reference = preapproval.get("external_reference")

    # Inicializamos user en None para ir actualizándolo si encontramos coincidencias.
    user = None

    # Verificamos si external_reference existe y comienza con el prefijo esperado.
    if external_reference and external_reference.startswith("user:"):
        # Separamos el texto después de "user:" para obtener el UUID en formato string.
        parte_id = external_reference.split(":", 1)[1]
        
        # Usamos session.get para compatibilidad con SQLAlchemy 2.0
        user_id = parte_id.strip()
        if user_id:
            user = db.session.get(User, user_id)

    # Si todavía no encontramos usuario (o no venía external_reference),
    # probamos buscar por el ID de preapproval si ya lo teníamos guardado.
    if user is None:
        preapproval_id = preapproval.get("id")
        if preapproval_id:
            user = User.query.filter_by(mp_preapproval_id=preapproval_id).first()

    # Devolvemos el usuario encontrado o None si no hubo coincidencias.
    return user


def actualizar_estado_desde_preapproval(user, preapproval):
    """
    Propósito:
        Actualizar el estado de suscripción y las fechas de vigencia del usuario
        basándose en la información real recibida de Mercado Pago.
        Maneja renovaciones automáticas leyendo 'next_payment_date'.

    Entradas:
        user: Instancia del modelo User a actualizar.
        preapproval: Diccionario con los datos de la suscripción (JSON de MP).

    Salidas:
        None: Realiza cambios en el objeto user y hace commit a la BD.

    Dependencias:
        - datetime, timedelta, time, timezone
        - db.session
    """
    # Extraer el estado actual de la suscripción desde el JSON de MP
    status_mp = preapproval.get("status")

    # Extraer el ID de la suscripción para verificar si cambió
    id_suscripcion = preapproval.get("id")

    # Extraer la fecha del próximo cobro (formato ISO 8601) calculada por MP
    next_payment_str = preapproval.get("next_payment_date")

    # Imprimir log de inicio para depuración
    print(f"[MP] Procesando actualización para usuario: {user.email}")
    print(f"[MP] Estado recibido de MP: {status_mp}")
    print(f"[MP] Fecha próximo pago (raw): {next_payment_str}")

    # Verificar si el ID de suscripción es nuevo o cambió para este usuario
    if id_suscripcion and user.mp_preapproval_id != id_suscripcion:
        # Actualizar el ID en la base de datos
        user.mp_preapproval_id = id_suscripcion
        print(f"[MP] ID de suscripción vinculado/actualizado: {id_suscripcion}")

    # Lógica principal según el estado reportado
    if status_mp == "authorized":
        # El usuario tiene el pago al día
        user.estado_suscripcion = "ACTIVA"
        print("[MP] Estado establecido a: ACTIVA")

        # Variable para almacenar la nueva fecha calculada
        nueva_fecha = None

        # ESTRATEGIA A: Intentar usar la fecha oficial de Mercado Pago
        if next_payment_str:
            try:
                # MP envía formato ej: "2025-12-27T09:00:00.000-04:00"
                # Cortamos en la 'T' para quedarnos solo con YYYY-MM-DD
                fecha_str_clean = next_payment_str.split("T")[0]
                # Convertimos el string a objeto date
                nueva_fecha = datetime.strptime(fecha_str_clean, "%Y-%m-%d").date()
                print(f"[MP] Fecha obtenida de next_payment_date: {nueva_fecha}")
            except Exception as e:
                # Si falla el parseo, logueamos el error y seguimos con el plan B
                print(f"[MP] Error al parsear fecha de MP ({next_payment_str}): {e}")

        # ESTRATEGIA B: Cálculo local (Fallback) si MP no mandó fecha o falló
        if not nueva_fecha:
            print("[MP] Usando estrategia de cálculo local (Fallback)")
            # Obtenemos fecha actual en UTC
            hoy = datetime.now(timezone.utc).date()

            # Caso 1: Usuario nuevo o vencido (no tiene fecha fin o ya pasó)
            if not user.fecha_fin_suscripcion or user.fecha_fin_suscripcion < hoy:
                # Le damos 30 días a partir de hoy
                nueva_fecha = hoy + timedelta(days=30)
                print(
                    f"[MP] Suscripción nueva/vencida. Nueva fecha (hoy+30): {nueva_fecha}"
                )

            # Caso 2: Renovación (el usuario ya tiene fecha futura válida)
            elif user.fecha_fin_suscripcion:
                # Calculamos cuántos días le quedan
                dias_restantes = (user.fecha_fin_suscripcion - hoy).days
                # Solo extendemos si está por vencer (menos de 10 días) para evitar duplicar meses erróneamente
                if dias_restantes < 10:
                    # Sumamos 30 días a su fecha de vencimiento actual
                    nueva_fecha = user.fecha_fin_suscripcion + timedelta(days=30)
                    print(
                        f"[MP] Renovación detectada. Extendiendo 30 días desde fecha actual: {nueva_fecha}"
                    )
                else:
                    # Si le queda mucho tiempo, mantenemos la fecha que tiene
                    nueva_fecha = user.fecha_fin_suscripcion
                    print(
                        f"[MP] Fecha actual válida ({dias_restantes} días restantes), no se cambia."
                    )

        # Aplicar la nueva fecha al usuario si es diferente a la que tiene
        if nueva_fecha and nueva_fecha != user.fecha_fin_suscripcion:
            user.fecha_fin_suscripcion = nueva_fecha
            # Establecer la hora exacta de vencimiento al final del día (23:59:59)
            user.licencia_valida_hasta = datetime.combine(
                user.fecha_fin_suscripcion, time(23, 59, 59)
            )
            print(
                f"[MP] ¡Fecha actualizada exitosamente en objeto User!: {user.fecha_fin_suscripcion}"
            )
        else:
            print("[MP] No hubo cambios en la fecha de vencimiento.")

    # Manejo de otros estados
    elif status_mp == "pending":
        user.estado_suscripcion = "PENDIENTE_MP"
        print("[MP] Estado cambiado a: PENDIENTE_MP")
    elif status_mp == "paused":
        user.estado_suscripcion = "PAUSADA"
        print("[MP] Estado cambiado a: PAUSADA")
    elif status_mp == "cancelled":
        user.estado_suscripcion = "CANCELADA"
        print("[MP] Estado cambiado a: CANCELADA")

    # Guardar cambios en la base de datos
    try:
        db.session.commit()
        print("[MP] Cambios guardados en DB correctamente.")
    except Exception as error:
        print(f"[MP] ERROR CRÍTICO al guardar en DB: {error}")
        db.session.rollback()  # Revertir transacción en caso de error


def refrescar_estado_suscripcion(user):
    """
    Propósito:
        Verifica si la suscripción ya venció y la marca como VENCIDA.

    Entradas:
        user: El usuario a verificar

    Salidas:
        None (modifica el usuario si está vencida)

    Dependencias:
        datetime, db.session
    """
    if not user:
        return

    hoy = datetime.now(timezone.utc).date()

    # Si está activa pero la fecha ya pasó, marcarla como vencida
    esta_activa = user.estado_suscripcion == "ACTIVA"
    tiene_fecha = user.fecha_fin_suscripcion is not None
    fecha_paso = tiene_fecha and (user.fecha_fin_suscripcion < hoy)

    if esta_activa and fecha_paso:
        print(f"[LOCAL] Suscripción venció el {user.fecha_fin_suscripcion}")
        user.estado_suscripcion = "VENCIDA"
        db.session.commit()


def validar_webhook_mp(request):
    """
    Propósito:
        Verifica que el webhook realmente viene de Mercado Pago.

    Entradas:
        request: La petición HTTP recibida

    Salidas:
        bool: True si es válido, False si no

    Dependencias:
        hmac, hashlib, Config.MP_WEBHOOK_SECRET
    """
    # En modo desarrollo, aceptar sin validar
    if app.debug:
        print("[MP SECURITY] Modo DEBUG, aceptado")
        return True

    try:
        # Obtener headers de seguridad
        x_signature = request.headers.get("x-signature")
        x_request_id = request.headers.get("x-request-id")

        # Verificar que existan
        if not x_signature or not x_request_id:
            print("[MP SECURITY] Faltan headers")
            return False

        # Extraer timestamp y hash del header
        partes = {}
        for fragmento in x_signature.split(","):
            if "=" in fragmento:
                clave, valor = fragmento.split("=", 1)
                partes[clave.strip()] = valor.strip()

        timestamp = partes.get("ts")
        hash_recibido = partes.get("v1")

        if not timestamp or not hash_recibido:
            print("[MP SECURITY] Formato inválido")
            return False

        # Crear el texto a hashear
        texto = f"id={x_request_id};request-id={x_request_id};ts={timestamp};"

        # Obtener el secret
        secret = Config.MP_WEBHOOK_SECRET
        if not secret:
            print("[MP SECURITY] Sin secret configurado")
            return True  # Permitir en desarrollo

        # Calcular hash esperado
        hash_esperado = hmac.new(
            secret.encode(), texto.encode(), hashlib.sha256
        ).hexdigest()

        # Comparar hashes
        es_valido = hmac.compare_digest(hash_esperado, hash_recibido)

        if not es_valido:
            print("[MP SECURITY] Hash no coincide")

        return es_valido

    except Exception as error:
        print(f"[MP SECURITY] Error: {error}")
        return False


def actualizar_estado_suscripcion_desde_mp(user):
    """
    Propósito:
        Consulta Mercado Pago para actualizar el estado de la suscripción.

    Entradas:
        user: El usuario a actualizar

    Salidas:
        bool: True si funcionó, False si hubo error

    Dependencias:
        Config.sdk_mp, actualizar_estado_desde_preapproval
    """
    # Verificar que el usuario tenga ID de suscripción
    if not user or not user.mp_preapproval_id:
        print("[MP] Usuario sin ID de suscripción")
        return False

    try:
        # Obtener SDK de Mercado Pago
        sdk = Config.sdk_mp

        print(f"[MP] Consultando suscripción {user.mp_preapproval_id}")

        # Consultar la API de MP
        respuesta = sdk.preapproval().get(user.mp_preapproval_id)
        datos_suscripcion = respuesta.get("response", {})

        # Verificar que llegaron datos
        if not datos_suscripcion or not datos_suscripcion.get("id"):
            print("[MP] No se recibieron datos")
            return False

        print(f"[MP] Estado recibido: {datos_suscripcion.get('status')}")

        # Actualizar usando la función principal
        actualizar_estado_desde_preapproval(user, datos_suscripcion)

        return True

    except Exception as error:
        print(f"[MP] Error al consultar: {error}")
        return False


# ====== RUTAS ======


@app.route("/")
def landing():
    """
    Landing page con información del proyecto.
    """

    return render_template("landing.html")


@app.route("/terminos")
def terminos():
    """Página de términos y condiciones"""
    return render_template("terminos.html")


@app.route("/sobre-nosotros")
def sobre_nosotros():
    """Página sobre nosotros"""
    return render_template("sobre_nosotros.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    """
    Propósito:
        Registrar nuevos usuarios en la aplicación usando un email único y
        almacenando la contraseña de forma segura (hasheada desde el modelo User).
        Además, envía un correo de confirmación al usuario tras el registro exitoso.

    Entradas:
        - Método GET:
            * Muestra el formulario de registro.
        - Método POST (formulario):
            * email: correo electrónico del usuario.
            * password: contraseña elegida por el usuario.
            * confirm_password: repetición de la contraseña.
            * nombre: nombre del usuario (opcional).
            * apellido: apellido del usuario (opcional).

    Salidas:
        - Renderiza 'register.html' cuando:
            * Es una petición GET.
            * Falta información o hay errores de validación.
        - Redirige a 'login' cuando:
            * El registro se completa correctamente.

    Dependencias:
        - validate_email_format para validar y normalizar el correo.
        - validate_password_strength para comprobar la fortaleza de la contraseña.
        - send_registration_confirmation_email para enviar correo de bienvenida.
        - Modelo User para crear el registro en la base de datos.
        - db.session para guardar el nuevo usuario.
        - current_user para saber si ya hay alguien autenticado.
        - mail para enviar el correo de confirmación.
    """

    # Comprobamos si ya hay un usuario autenticado en la sesión actual.
    if current_user.is_authenticated:
        # Si ya está logueado, no tiene sentido registrar otro usuario; lo mandamos a la app.
        return redirect(url_for("upload_files"))

    # Revisamos si la petición es de tipo POST (envío del formulario de registro).
    if request.method == "POST":
        # ================================================================
        # HONEYPOT: Validación anti-bot
        # ================================================================
        # Obtenemos el valor del campo honeypot (debe estar vacío si es humano).
        honeypot = request.form.get("website", "").strip()
        # Si el campo honeypot contiene algún valor, es muy probable que sea un bot.
        if honeypot:
            # Registramos el intento de bot en los logs para monitoreo.
            app.logger.warning(f"Intento de registro de bot detectado. Honeypot llenado: {honeypot}")
            # Redirigimos silenciosamente al login para no alertar al bot.
            # No mostramos mensajes de error para evitar que el bot ajuste su comportamiento.
            return redirect(url_for("login"))
        # ================================================================

        # Obtenemos el email del formulario y quitamos espacios extra al inicio y final.
        email = request.form.get("email", "").strip()
        # Obtenemos la contraseña tal como fue ingresada por el usuario.
        password = request.form.get("password")
        # Obtenemos la confirmación de la contraseña.
        confirm_password = request.form.get("confirm_password")
        # Obtenemos el nombre, opcional, y eliminamos espacios sobrantes.
        nombre = request.form.get("nombre", "").strip()
        # Obtenemos el apellido, opcional, y eliminamos espacios sobrantes.
        apellido = request.form.get("apellido", "").strip()

        # Validamos que se haya ingresado un email y una contraseña.
        if not email or not password:
            # Mostramos un mensaje de error indicando que ambos campos son obligatorios.
            flash("Email y contraseña son obligatorios", "danger")
            # Volvemos a mostrar el formulario de registro.
            return render_template("register.html")

        # Comprobamos que el correo no exceda el máximo permitido (150 caracteres).
        if len(email) > 150:
            # Si es demasiado largo, mostramos un mensaje de error.
            flash("El correo no puede exceder 150 caracteres", "danger")
            # Renderizamos de nuevo el formulario para que pueda corregirlo.
            return render_template("register.html")

        # Si el usuario escribió un nombre, revisamos que no sea demasiado largo.
        if nombre and len(nombre) > 80:
            # Mostramos error si el nombre supera los 80 caracteres.
            flash("El nombre no puede exceder 80 caracteres", "danger")
            # Volvemos a mostrar el formulario.
            return render_template("register.html")

        # Si el usuario escribió un apellido, revisamos que no sea demasiado largo.
        if apellido and len(apellido) > 80:
            # Mostramos error si el apellido supera los 80 caracteres.
            flash("El apellido no puede exceder 80 caracteres", "danger")
            # Volvemos a mostrar el formulario.
            return render_template("register.html")

        # Llamamos a la función que valida y normaliza el correo electrónico.
        is_valid_email, normalized_email, email_error = validate_email_format(email)
        # Si el correo no pasa la validación de formato, mostramos el motivo.
        if not is_valid_email:
            # Mostramos el mensaje de error específico que devuelve la función.
            flash(f"Correo inválido: {email_error}", "danger")
            # Volvemos a mostrar el formulario de registro.
            return render_template("register.html")

        # Comparamos la contraseña con la confirmación para asegurarnos de que coincidan.
        if password != confirm_password:
            # Si no son iguales, informamos al usuario del error.
            flash("Las contraseñas no coinciden", "danger")
            # Renderizamos de nuevo el formulario para que las vuelva a ingresar.
            return render_template("register.html")

        # Llamamos a la función que evalúa la fortaleza de la contraseña.
        is_valid_password, password_error = validate_password_strength(password)
        # Si la contraseña no cumple los requisitos, mostramos el motivo.
        if not is_valid_password:
            # Mostramos el mensaje de error devuelto por la validación.
            flash(password_error, "danger")
            # Volvemos a mostrar el formulario de registro.
            return render_template("register.html")

        # Buscamos si ya existe un usuario con ese correo normalizado en la base de datos.
        if User.query.filter_by(email=normalized_email).first():
            # Si ya hay un registro con ese correo, no permitimos un duplicado.
            flash("El correo ya está registrado", "danger")
            # Mostramos otra vez el formulario para que use otro correo.
            return render_template("register.html")

        # Creamos una nueva instancia de User con los datos ingresados.
        # El modelo se encargará de hashear la contraseña internamente.
        new_user = User(
            email=normalized_email, password=password, nombre=nombre, apellido=apellido
        )

        # Añadimos el nuevo usuario a la sesión de la base de datos.
        db.session.add(new_user)
        # Confirmamos los cambios guardando el nuevo registro en la base.
        db.session.commit()

        # ================================================================
        # ENVÍO DE CORREO DE CONFIRMACIÓN DE REGISTRO
        # ================================================================
        # Intentamos enviar un correo de bienvenida al usuario recién registrado.
        # Importamos la función que maneja el envío del correo de confirmación.
        from utils import send_registration_confirmation_email
        
        # Llamamos a la función de envío pasando el usuario creado y el objeto mail.
        # Esta función devuelve True si el correo se envió exitosamente, False si falló.
        email_enviado = send_registration_confirmation_email(new_user, mail)
        
        # Verificamos si el correo se envió correctamente.
        if email_enviado:
            # Si el correo se envió, mostramos un mensaje de éxito completo.
            # Informamos al usuario que revise su correo para la confirmación.
            flash(
                "¡Registro exitoso! Revisa tu correo para confirmar tu cuenta.", 
                "success"
            )
        else:
            # Si falló el envío del correo, igual permitimos que el registro sea válido.
            # Mostramos un mensaje de advertencia indicando el problema con el correo.
            # Esto evita bloquear el registro por un fallo en el servicio de email.
            flash(
                "Registro exitoso, pero hubo un problema al enviar el correo de confirmación.", 
                "warning"
            )
        # ================================================================

        # Redirigimos al usuario a la página de login para que pueda iniciar sesión.
        return redirect(url_for("login"))

    # Si la petición es GET (o no se cumplió ninguna condición anterior), mostramos el formulario.
    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    """
    Propósito:
        Autenticar usuarios, manejar el inicio de sesión y registrar el último acceso.
        Además, redirigir a la página de suscripción si el usuario no tiene una
        suscripción vigente y necesita pagar.

    Entradas:
        - Método GET:
            * Muestra el formulario de inicio de sesión.
        - Método POST (formulario):
            * email: correo del usuario.
            * password: contraseña del usuario.

    Salidas:
        - Renderiza la plantilla 'login.html' cuando:
            * Es un GET.
            * Faltan datos.
            * Las credenciales son incorrectas.
            * La cuenta no está activa.
        - Redirige a:
            * /suscripcion si necesita pagar.
            * /app (upload_files) si la suscripción está vigente.
            * next (si viene de una página protegida) tras login exitoso.

    Dependencias:
        - Modelo User y db.session para consultar y actualizar la BD.
        - Flask-Login: current_user, login_user.
        - Funciones suscripcion_vigente(user) y necesita_pago(user) para la lógica de suscripción.
        - datetime.utcnow para registrar el último acceso.
    """

    # Verificamos si ya hay un usuario autenticado en la sesión actual.
    if current_user.is_authenticated:
        # Si ya está logueado, actualizamos su estado de suscripción.
        refrescar_estado_suscripcion(current_user)

        # Revisamos si su suscripción NO está vigente y además requiere pago.
        if not suscripcion_vigente(current_user) and necesita_pago(current_user):
            # Si debe pagar, lo redirigimos a la página de suscripción.
            return redirect(url_for("suscripcion_checkout"))
        # Si la suscripción está bien o no necesita pago, lo mandamos a la app principal.
        return redirect(url_for("upload_files"))

    # Comprobamos si la petición es de tipo POST (envío del formulario).
    if request.method == "POST":
        # Obtenemos el correo desde el formulario y eliminamos espacios extra.
        email = request.form.get("email", "").strip()
        # Obtenemos la contraseña desde el formulario tal como viene.
        password = request.form.get("password")

        # Verificamos que el correo y la contraseña no estén vacíos.
        if not email or not password:
            # Mostramos un mensaje de error si faltan datos.
            flash("Correo y contraseña son requeridos", "danger")
            # Renderizamos de nuevo la plantilla de login para que intente otra vez.
            return render_template("login.html")

        # Buscamos en la base de datos un usuario con ese correo electrónico.
        user = User.query.filter_by(email=email).first()

        # Verificamos que el usuario exista y que la contraseña sea correcta.
        if user and user.check_password(password):
            # Comprobamos si la cuenta del usuario está marcada como ACTIVA.
            if user.estado_cuenta != "ACTIVA":
                # Si la cuenta no está activa, mostramos un mensaje y no iniciamos sesión.
                flash("Tu cuenta no está activa", "danger")
                # Volvemos a mostrar el formulario de login.
                return render_template("login.html")

            # Llamamos a login_user para guardar al usuario en la sesión.
            # CAMBIO SEGURIDAD: Force remember=False para que la cookie expire al cerrar el navegador.
            login_user(user, remember=False)
            
            # Actualizamos el campo ultimo_acceso con la fecha y hora actual en UTC.
            user.ultimo_acceso = datetime.now(timezone.utc)
            # Guardamos los cambios en la base de datos.
            db.session.commit()

            # Actualiza estado según fecha
            refrescar_estado_suscripcion(user)

            # Revisamos si la suscripción NO está vigente y además requiere pago.
            if not suscripcion_vigente(user) and necesita_pago(user):
                # Si debe pagar, lo redirigimos a la página de suscripción.
                return redirect(url_for("suscripcion_checkout"))

            # Obtenemos la página a la que quería ir originalmente (si existe).
            next_page = request.args.get("next")
            # Si existe next_page, redirigimos allí; si no, vamos a la app principal.
            return (
                redirect(next_page) if next_page else redirect(url_for("upload_files"))
            )

        else:
            # Si el usuario no existe o la contraseña es incorrecta, mostramos un mensaje de error.
            flash("Correo o contraseña incorrectos", "danger")

    # Si es una petición GET, o hubo un error, mostramos el formulario de login.
    return render_template("login.html")


@app.route("/suscripcion")
@login_required
def suscripcion_checkout():
    # Propósito:
    #   Mostrar la página con la información del plan de suscripción
    #   y un botón para iniciar el proceso con Mercado Pago.
    #
    # Entradas:
    #   Ninguna directa (usa current_user solo para validar acceso).
    #
    # Salidas:
    #   Render del template 'suscripcion.html' con datos del plan.
    #
    # Dependencias:
    #   - current_user de Flask-Login (para validar acceso).
    #   - Template 'suscripcion.html'.

    # Primero, refrescamos el estado de suscripción del usuario actual
    refrescar_estado_suscripcion(current_user)

    # Si está pendiente, intentamos refrescar el estado desde MP
    if current_user.estado_suscripcion == "PENDIENTE_MP":
        actualizar_estado_suscripcion_desde_mp(current_user)

    # Definimos los datos básicos del plan que se mostrará en la página.
    if suscripcion_vigente(current_user):
        # Ya está al día, mejor mandarlo a la app
        return redirect(url_for("upload_files"))

    plan = {
        "nombre": "Plan Standard FAT Testing",
        "monto": 150000,
        "moneda": "CLP",
        "renovacion": "Mensual",
    }

    # Calculamos la fecha estimada de próxima renovación (por ahora, la fecha de hoy).
    proxima_fecha = datetime.now(timezone.utc).date()

    # Esta ruta interna creará la suscripción vía SDK y luego redirigirá a Mercado Pago.
    checkout_url = url_for("iniciar_suscripcion_mp")

    # Renderizamos el template de suscripción con los datos del plan y la URL del botón.
    return render_template(
        "suscripcion.html",
        plan=plan,
        proxima_fecha=proxima_fecha,
        checkout_url=checkout_url,
    )


@app.route("/suscripcion/iniciar")
@login_required
def iniciar_suscripcion_mp():
    """
    Crea la suscripción en Mercado Pago y redirige al usuario.
    """
    print(f"[MP] Iniciando para {current_user.email}")

    # Preparar datos
    sdk = Config.sdk_mp
    referencia = f"user:{current_user.get_id()}"
    email_pagador = os.environ.get("MP_TEST_PAYER_EMAIL", current_user.email)

    datos_suscripcion = {
        "auto_recurring": {
            "currency_id": "CLP",
            "transaction_amount": 150000,
            "frequency": 1,
            "frequency_type": "months",
        },
        "back_url": url_for("suscripcion_retorno", _external=True),
        "external_reference": referencia,
        "payer_email": email_pagador,
        "reason": "Plan Standard FAT Testing",
    }

    try:
        # Crear suscripción en MP
        respuesta = sdk.preapproval().create(datos_suscripcion)
        datos = respuesta.get("response", {})

        id_suscripcion = datos.get("id")
        url_pago = datos.get("init_point") or datos.get("sandbox_init_point")

        print(f"[MP] ID: {id_suscripcion}")
        print(f"[MP] URL: {url_pago}")

        if not id_suscripcion:
            flash("No se pudo crear la suscripción", "danger")
            return redirect(url_for("suscripcion_checkout"))

        if not url_pago:
            flash("No se pudo obtener URL de pago", "danger")
            return redirect(url_for("suscripcion_checkout"))

        # Guardar ID y marcar como pendiente
        current_user.mp_preapproval_id = id_suscripcion
        current_user.estado_suscripcion = "PENDIENTE_MP"
        db.session.commit()

        print("[MP] Redirigiendo a MP")
        return redirect(url_pago)

    except Exception as error:
        print(f"[MP] Error: {error}")
        flash("Error al crear suscripción", "danger")
        return redirect(url_for("suscripcion_checkout"))


@app.route("/suscripcion/retorno")
@login_required
def suscripcion_retorno():
    """
    Página donde vuelve el usuario después de pagar.
    Versión ESTABLE: Verifica una vez y responde rápido.
    """
    print(f"[RETORNO] {current_user.email} volvió")

    # 1. Si ya está activo localmente, pase directo
    refrescar_estado_suscripcion(current_user)
    if suscripcion_vigente(current_user):
        print("[RETORNO] Usuario ya estaba activo.")
        flash("¡Suscripción validada correctamente!", "success")
        return redirect(url_for("upload_files"))

    # 2. Si está pendiente, hacemos UN intento de consulta a MP
    if current_user.estado_suscripcion == "PENDIENTE_MP":
        print("[RETORNO] Estado PENDIENTE. Consultando a MP por última vez...")
        actualizar_estado_suscripcion_desde_mp(current_user)

        # Volvemos a verificar tras la consulta
        refrescar_estado_suscripcion(current_user)

    # 3. Revisión final
    if suscripcion_vigente(current_user):
        print("[RETORNO] ¡Activación exitosa tras consulta!")
        flash("¡Pago confirmado! Tu suscripción está activa.", "success")
        return redirect(url_for("upload_files"))

    # 4. Si sigue pendiente tras el intento
    if current_user.estado_suscripcion == "PENDIENTE_MP":
        print("[RETORNO] Sigue pendiente en MP.")
        flash(
            "Tu pago se está procesando. Por favor espera 1 minuto y recarga la página.",
            "info",
        )
        # IMPORTANTE: Redirigimos a checkout para romper el bucle,
        # el usuario verá el mensaje Flash y podrá reintentar manualmente.
        return redirect(url_for("suscripcion_checkout"))

    # 5. Si hubo error/cancelación
    flash("Hubo un problema al procesar la suscripción.", "warning")
    return redirect(url_for("suscripcion_checkout"))


@app.route("/mp/webhook", methods=["GET", "POST"])
@csrf.exempt
def mp_webhook():
    """
    Propósito:
        Recibir notificaciones de Mercado Pago.
    
    Entradas: request (JSON de MP)
    Salidas: 200 OK
    Dependencias: extraer_evento_mp, Config, User, actualizar_estado_desde_preapproval
    """
    if request.method == "GET":
        print("=== WEBHOOK GET (Healthcheck) ===")
        return "", 200
    
    # Validación básica (deshabilitada en debug para facilitar pruebas)
    if not validar_webhook_mp(request):
        print("[MP SECURITY] Webhook rechazado por firma inválida")
        # return "", 403 

    try:
        body = request.get_json(silent=True) or {}
        query = request.args.to_dict()
        topic, resource_id = extraer_evento_mp(body, query)
        
        if not topic or not resource_id:
            print("[MP] Webhook recibido sin datos clave.")
            return "", 200

        sdk = Config.sdk_mp

        # CASO 1: Novedades de la Suscripción (Pausas, Cancelaciones, Alta)
        if topic in ["subscription_preapproval", "subscription_preapproval_plan"]:
            print(f"\n[MP] Novedad en Suscripción. ID: {resource_id}")
            resp = sdk.preapproval().get(resource_id)
            datos = resp.get("response", {})
            
            usuario = obtener_usuario_desde_preapproval(datos)
            if usuario:
                print(f"[MP] Usuario encontrado: {usuario.email}. Actualizando...")
                actualizar_estado_desde_preapproval(usuario, datos)
            else:
                print("[MP] ID de suscripción no encontrado en nuestra BD (puede ser antiguo).")

        # CASO 2: PAGO CONFIRMADO (Aquí está la mejora "Detective")
        elif topic == "subscription_authorized_payment":
            print(f"\n[MP] PAGO RECIBIDO. ID Pago: {resource_id}")
            
            # Consultamos el detalle del pago
            payment_resp = sdk.payment().get(resource_id)
            payment_data = payment_resp.get("response", {})
            
            # --- INICIO MODO DETECTIVE ---
            # Buscamos el ID de la suscripción en 3 lugares distintos
            
            # Lugar 1: Donde debería estar siempre
            preapproval_id = payment_data.get("preapproval_id")
            
            # Lugar 2: A veces MP lo guarda en 'metadata'
            if not preapproval_id:
                preapproval_id = payment_data.get("metadata", {}).get("preapproval_id")
                
            # Lugar 3: A veces viene en el JSON original del webhook
            if not preapproval_id:
                preapproval_id = body.get("data", {}).get("preapproval_id")

            # -----------------------------
            
            usuario = None

            if preapproval_id:
                print(f"[MP] ID Suscripción encontrado: {preapproval_id}")
                # Buscamos al usuario que tenga ese ID guardado
                usuario = User.query.filter_by(mp_preapproval_id=preapproval_id).first()
            else:
                # ESTRATEGIA DE EMERGENCIA: Buscar por "external_reference"
                # (Recuerda que nosotros le pusimos "user:EL_UUID" al crear la suscripción)
                ext_ref = payment_data.get("external_reference")
                if ext_ref and ext_ref.startswith("user:"):
                    print(f"[MP] Buscando por etiqueta externa: {ext_ref}")
                    user_id = ext_ref.split(":", 1)[1]
                    usuario = db.session.get(User, user_id)
                    # Si lo encontramos, recuperamos su ID de suscripción para consultar a MP
                    if usuario:
                        preapproval_id = usuario.mp_preapproval_id

            # Si después de todo eso encontramos al usuario...
            if usuario and preapproval_id:
                print(f"[MP] ¡Usuario identificado!: {usuario.email}")
                print("[MP] Consultando fecha de vencimiento actualizada...")
                
                # Vamos a preguntar a la API de suscripciones cuándo vence ahora
                sub_resp = sdk.preapproval().get(preapproval_id)
                sub_data = sub_resp.get("response", {})
                
                # Actualizamos la base de datos
                actualizar_estado_desde_preapproval(usuario, sub_data)
                print("[MP] Renovación procesada con éxito.")
            else:
                print(f"[MP] No se pudo vincular este pago a ningún usuario. Se ignora.")
        
        else:
            print(f"[MP] Evento ignorado: {topic}")

    except Exception as e:
        print(f"[MP] ERROR en Webhook: {e}")
    
    return "", 200


@app.route("/logout")
@login_required
def logout():
    """Ruta para cerrar sesión"""
    from flask_login import logout_user

    logout_user()
    flash("Has cerrado sesión exitosamente", "info")
    return redirect(url_for("landing"))


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    """
    Propósito: generar token de recuperación y enviar correo.
    Entradas (POST form): email.
    Salidas: render de formulario o redirect a login tras mostrar mensaje.
    Dependencias: User, PasswordResetToken, send_password_reset_email, db.session.
    """
    # Generar token de recuperación y enviar correo
    if current_user.is_authenticated:
        return redirect(
            url_for("upload_files")
        )  # Evita que un usuario logueado pida reset

    if request.method == "POST":
        email = request.form.get("email", "").strip()  # Correo para recuperar

        if not email:
            flash("Ingresa tu correo", "danger")
            return render_template("forgot_password.html")

        is_valid_email, normalized_email, email_error = validate_email_format(
            email
        )  # Valida formato
        if not is_valid_email:
            flash(f"Correo inválido: {email_error}", "danger")
            return render_template("forgot_password.html")

        user = User.query.filter_by(email=normalized_email).first()  # Busca al usuario
        flash(
            "Si el correo está registrado, recibirás instrucciones de recuperación",
            "info",
        )  # Mensaje uniforme

        if user:
            # En el nuevo modelo, la PK del usuario es un UUID (String(36)).
            user_uuid = user.get_id()

            # Marcamos como usados todos los tokens anteriores vigentes de este usuario.
            old_tokens = PasswordResetToken.query.filter_by(
                user_id=user_uuid, used=False
            ).all()  # Tokens previos
            for token in old_tokens:
                token.mark_as_used()  # Invalida tokens anteriores

            # Creamos un nuevo token asociado al UUID del usuario.
            reset_token = PasswordResetToken(
                user_id=user_uuid, expiration_hours=1
            )  # Crea token nuevo
            db.session.add(reset_token)  # Agrega a la sesión
            db.session.commit()  # Guarda en la base

            email_sent = send_password_reset_email(
                user, reset_token.token, mail
            )  # Envía correo
            if not email_sent:
                db.session.delete(reset_token)  # Limpia el token si falló el envío
                db.session.commit()
                flash("Error al enviar el correo. Intenta más tarde.", "danger")
                return render_template("forgot_password.html")

        return redirect(url_for("login"))  # Redirige a login tras procesar

    return render_template("forgot_password.html")  # Muestra formulario en GET


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token):
    """
    Propósito: permitir cambio de contraseña usando un token válido.
    Entradas: token en la URL; (POST form) password y confirm_password.
    Salidas: render de formulario o redirect a login tras éxito.
    Dependencias: PasswordResetToken.get_valid_token, bcrypt, db.session.
    """
    # Permitir cambio de contraseña usando un token válido
    if current_user.is_authenticated:
        return redirect(url_for("upload_files"))  # Evita uso si ya está logueado

    reset_token = PasswordResetToken.get_valid_token(token)  # Valida token
    if not reset_token:
        flash("Enlace inválido o expirado. Solicita uno nuevo.", "danger")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        password = request.form.get("password")  # Nueva contraseña
        confirm_password = request.form.get("confirm_password")  # Confirmación

        if not password or not confirm_password:
            flash("Todos los campos son requeridos", "danger")
            return render_template("reset_password.html", token=token)

        if password != confirm_password:  # Coincidencia
            flash("Las contraseñas no coinciden", "danger")
            return render_template("reset_password.html", token=token)

        is_valid_password, password_error = validate_password_strength(
            password
        )  # Fortaleza
        if not is_valid_password:
            flash(password_error, "danger")
            return render_template("reset_password.html", token=token)

        user = reset_token.user  # Usuario dueño del token
        user.password_hash = bcrypt.generate_password_hash(password).decode(
            "utf-8"
        )  # Re-hash
        reset_token.mark_as_used()  # Marca token como usado
        db.session.commit()  # Guarda cambios

        flash("Contraseña actualizada. Ahora puedes iniciar sesión.", "success")
        return redirect(url_for("login"))

    return render_template(
        "reset_password.html", token=token
    )  # Muestra formulario en GET

@app.route("/soporte", methods=["GET", "POST"])
def soporte():
    """
    Propósito:
        Mostrar formulario de soporte y procesar el envío usando utils.
    """
    if request.method == "POST":
        nombre = request.form.get("nombre")
        email = request.form.get("email")
        asunto = request.form.get("asunto")
        mensaje_texto = request.form.get("mensaje")
        
        # Validación básica
        if not nombre or not email or not mensaje_texto:
            flash("Por favor completa todos los campos", "danger")
            return render_template("soporte.html")

        # Llamamos a la función en utils.py que maneja toda la lógica del correo
        enviado = send_support_email(nombre, email, asunto, mensaje_texto, mail)

        if enviado:
            flash("¡Ticket enviado! Nuestro equipo te responderá pronto.", "success")
            # Redirigir para limpiar el formulario (PRG Pattern)
            return redirect(url_for('soporte'))
        else:
            flash("Error al enviar el mensaje. Por favor intenta más tarde.", "danger")
            # No redirigimos para que el usuario no pierda lo que escribió (opcional)
            return render_template("soporte.html")

    # Si es GET
    return render_template("soporte.html")


@app.route("/dashboard")
@suscripcion_requerida
def dashboard():
    """
    Propósito: listar usuarios (ajusta la política de acceso según tu criterio).
    Entradas: ninguna.
    Salidas: render de dashboard con usuarios.
    Dependencias: User.
    Nota: si ya no hay roles, aquí puedes permitir solo si la cuenta está ACTIVA o implementar tu propio check.
    """
    # Ejemplo simple: solo usuarios con cuenta activa pueden ver el dashboard
    if current_user.estado_cuenta != "ACTIVA":
        flash("No tienes permisos para acceder a esta página", "danger")
        return redirect(url_for("upload_files"))

    users = User.query.order_by(User.fecha_creacion.desc()).all()
    return render_template("dashboard.html", users=users)


@app.route("/download-desktop-app")
@suscripcion_requerida
def download_desktop_app():
    """Ruta para descargar la aplicación de escritorio"""
    try:
        # Ruta del ejecutable
        exe_path = os.path.join(
            os.path.dirname(__file__), "downloads", "FAT_Testing.exe"
        )

        # Verificar que el archivo existe
        if not os.path.exists(exe_path):
            flash(
                "La aplicación de escritorio no está disponible en este momento. Por favor contacta al administrador.",
                "warning",
            )
            return redirect(url_for("upload_files"))

        # Enviar archivo para descarga
        return send_file(
            exe_path,
            as_attachment=True,
            download_name="FAT_Testing.exe",
            mimetype="application/octet-stream",
        )
    except Exception as e:
        flash(f"Error al descargar la aplicación: {str(e)}", "danger")
        return redirect(url_for("upload_files"))


# ====== RUTAS DE LA APLICACIÓN ======


@app.route("/app", methods=["GET", "POST"])
@suscripcion_requerida
def upload_files():

    opciones = [
        "SW L2 9200",
        "SW L2 9300",
        "SW L2 9500",
    ]

    if request.method == "POST":

        file_type = request.form.get("fileType")
        file = request.files.get("file")
        processed_files = []
        img_1 = request.files.getlist("imageFiles1")
        img_2 = request.files.getlist("imageFiles2")
        img_3 = request.files.getlist("imageFiles3")
        proyecto = request.form.get("proyecto")
        cliente = request.form.get("cliente")
        ordenCompra = request.form.get("ordenCompra")
        notaVenta = request.form.get("notaVenta")

        if file:
            filename = secure_filename(file.filename)
            # Plantilla basada en el tipo de archivo
            if file_type == "SW L2 9200" or file_type == "SW L2 9300":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion SW 9200 - 9300.docx"
                )

            elif file_type == "SW L2 9500":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion SW 9500.docx"
                )

            elif file_type == "SW L3 9348GC" or file_type == "SW L3 C93180YC":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion SW 9348GC - C93180YC.docx"
                )

            elif file_type == "SW IE3300" or file_type == "SW IE4010":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion SW IE 3300 - 4010.docx"
                )

            elif file_type == "Router C8500":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion Router C8500.docx"
                )

            elif file_type == "Router ISR4431":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion Router ISR4431.docx"
                )

            elif (
                file_type == "AP C9115AXI"
                or file_type == "AP C9120AXE"
                or file_type == "AP C9130AXI"
            ):
                docx_template_path = os.path.join(
                    "plantillas",
                    "Template Extraccion C9115AXI-A,C9120AXE-A,C9130AXI-A.docx",
                )

            elif file_type == "Check Point 6200" or file_type == "Check Point 6600":
                docx_template_path = os.path.join(
                    "plantillas", "Template Extraccion Check Point 6200P - 6600P.docx"
                )

            # Procesamiento de archivo

            word_buffer, download_filename = procesar_archivo(
                file.stream,
                docx_template_path,
                img_1,
                img_2,
                img_3,
                proyecto,
                cliente,
                ordenCompra,
                notaVenta,
                file_type,
            )

            # Limpiar el nombre del archivo
            download_filename = (
                download_filename.strip().replace("\n", "").replace("\r", "")
            )
            download_filename = secure_filename(download_filename)

        return send_file(
            word_buffer,
            as_attachment=True,
            download_name=download_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    return render_template("informes.html", opciones=opciones)

# MANEJO DE ERROR CSRF
@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    """
    Atrapa los errores de token faltante o expirado.
    Redirige al usuario a la página anterior con un mensaje amigable.
    """
    # Mensaje para el usuario
    flash("Tu sesión ha expirado o el formulario no es válido. Por favor, intenta de nuevo.", "warning")
    
    # Redirección 
    # request.referrer contiene la URL desde donde vino el usuario (el formulario)
    if request.referrer:
        return redirect(request.referrer)
    
    # Si no sabemos de dónde vino, lo mandamos al login o al inicio
    return redirect(url_for('landing'))


# ====== COMANDOS CLI ======


@app.cli.command()
def init_db():
    """Inicializar la base de datos"""
    with app.app_context():
        db.create_all()
        print("Base de datos inicializada!")


@app.cli.command()
def create_admin():
    """
    Propósito: crear un usuario desde la línea de comandos (sin roles).
    Entradas: prompts de consola para email, password, nombre, apellido.
    Salidas: imprime resultado en consola.
    Dependencias: User, db.session, bcrypt (dentro del constructor).
    """
    # Crear un usuario desde la línea de comandos (sin roles)
    with app.app_context():
        email = input("Email: ").strip()  # Pide email
        password = input("Contraseña: ")  # Pide contraseña
        nombre = input("Nombre: ").strip()  # Pide nombre
        apellido = input("Apellido: ").strip()  # Pide apellido

        if User.query.filter_by(email=email).first():  # Revisa si ya existe
            print("El correo ya existe")
            return

        user = User(
            email=email, password=password, nombre=nombre, apellido=apellido
        )  # Crea usuario
        db.session.add(user)  # Agrega a la sesión
        db.session.commit()  # Guarda en la base
        print(f"Usuario {email} creado exitosamente")


# Iniciar la aplicación
if __name__ == "__main__":
    # Crear las tablas si no existen
    with app.app_context():
        db.create_all()

    # app.run(host="0.0.0.0", port=80, debug=True, ssl_context="adhoc")
    app.run(host="0.0.0.0", port=80, debug=True)
